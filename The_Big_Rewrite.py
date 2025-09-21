import os
from docx import Document
import tkinter as tk
import pyperclip
from tkinter import filedialog, messagebox
from docx.shared import RGBColor
class DocxReplacerApp:
    def __init__(self, root):
        self.root = root
        root.title("Замена текста с сохранением форматирования")
        root.geometry("500x350")
        
        self.folder_path = tk.StringVar()
        self.old_text = tk.StringVar()
        self.new_text = tk.StringVar()
        
        self.create_widgets()
        self.add_developer_label()
        self.setup_clipboard()  # Настраиваем обработку буфера обмена
    
    def setup_clipboard(self):
        # Добавляем горячую клавишу Ctrl+V для всего окна
        self.root.bind('<Control-v>', self.paste_from_clipboard)
        self.root.bind('<Control-V>', self.paste_from_clipboard)  # Для Caps Lock
    
    def paste_from_clipboard(self, event=None):
        try:
            # Получаем текст из буфера обмена
            clipboard_text = pyperclip.paste()
            if not clipboard_text:
                return
            
            # Определяем, какое поле сейчас в фокусе
            focused_widget = self.root.focus_get()
            
            if isinstance(focused_widget, tk.Entry):
                # Вставляем текст в активное поле ввода
                focused_widget.insert(tk.INSERT, clipboard_text)
        except Exception as e:
            print(f"Ошибка при вставке из буфера обмена: {str(e)}")

    def add_developer_label(self):
        footer_frame = tk.Frame(self.root)
        footer_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=5) 
            
        developer_label = tk.Label(
            footer_frame, 
            text="Разработчик: ОАТиМ Егор Власов",
            fg="gray",
            font=("Arial", 8)
        )
        developer_label.pack(side=tk.RIGHT)

    def create_widgets(self):
        tk.Label(self.root, text="Папка с документами:").pack(pady=(10, 0))
        
        path_frame = tk.Frame(self.root)
        path_frame.pack(fill=tk.X, padx=10)
        self.folder_entry = tk.Entry(path_frame, textvariable=self.folder_path, width=40)
        self.folder_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Добавляем контекстное меню для вставки
        self.setup_entry_context_menu(self.folder_entry)
        
        tk.Button(path_frame, text="Обзор", command=self.browse_folder).pack(side=tk.RIGHT)
        
        tk.Label(self.root, text="Текст для поиска:").pack(pady=(10, 0))
        self.old_text_entry = tk.Entry(self.root, textvariable=self.old_text, width=500)
        self.old_text_entry.pack()
        self.setup_entry_context_menu(self.old_text_entry)
        
        tk.Label(self.root, text="Заменить на:").pack(pady=(10, 0))
        self.new_text_entry = tk.Entry(self.root, textvariable=self.new_text, width=50)
        self.new_text_entry.pack()
        self.setup_entry_context_menu(self.new_text_entry)
        
        tk.Button(self.root, text="Выполнить замену", command=self.run_replacement, 
                bg="green", fg="white").pack(pady=20)
        
        self.status_label = tk.Label(self.root, text="", fg="blue")
        self.status_label.pack()
    
    def setup_entry_context_menu(self, entry_widget):
        # Создаем контекстное меню для поля ввода
        context_menu = tk.Menu(entry_widget, tearoff=0)
        context_menu.add_command(label="Вставить", command=lambda: self.paste_to_entry(entry_widget))
        
        # Привязываем меню к правой кнопке мыши
        entry_widget.bind("<Button-3>", lambda e: context_menu.tk_popup(e.x_root, e.y_root))
    
    def paste_to_entry(self, entry_widget):
        try:
            # Вставляем текст в указанное поле
            clipboard_text = pyperclip.paste()
            if clipboard_text:
                entry_widget.insert(tk.INSERT, clipboard_text)
        except Exception as e:
            print(f"Ошибка при вставке: {str(e)}")
    
    def browse_folder(self):
        folder_selected = filedialog.askdirectory()
        if folder_selected:
            self.folder_path.set(folder_selected)
    
    def run_replacement(self):
        folder_path = self.folder_path.get()
        old_text = self.old_text.get()
        new_text = self.new_text.get()
        
        if not all([folder_path, old_text]):
            messagebox.showerror("Ошибка", "Укажите папку и текст для поиска!")
            return
        
        try:
            total_files, total_replacements = self.replace_text_in_docx(folder_path, old_text, new_text)
            messagebox.showinfo("Готово", 
                            f"Обработка завершена!\n\n"
                            f"Файлов обработано: {total_files}\n"
                            f"Всего замен: {total_replacements}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка:\n{str(e)}")
    
    def replace_text_in_paragraph(self, paragraph, old_text, new_text):
        """Улучшенная замена текста с обработкой длинных предложений"""
        full_paragraph_text = paragraph.text
        if old_text not in full_paragraph_text:
            return 0
        
        # Сохраняем форматирование первого run
        first_run = paragraph.runs[0] if paragraph.runs else None
        if first_run:
            font = first_run.font
            # formatting = {
            #     'bold': font.bold,
            #     'italic': False,
            #     'underline': font.underline,
            #     'color':  RGBColor(0, 0, 0),
            #     'size': font.size,
            #     'name': font.name
            # }
            formatting = {
                'bold': font.bold,
                'italic': font.italic,
                'underline': font.underline,
                'color': font.color.rgb if font.color else None,
                'size': font.size,
                'name': font.name
            }
        
        # Полная замена во всем абзаце
        new_text_content = full_paragraph_text.replace(old_text, new_text)
        
        # Очищаем все runs
        for run in paragraph.runs:
            run.text = ""
        
        # Добавляем новый текст с сохранением форматирования
        if paragraph.runs:
            paragraph.runs[0].text = new_text_content
            if first_run:
                for run in paragraph.runs:
                    run.font.bold = formatting['bold']
                    run.font.italic = formatting['italic']
                    run.font.underline = formatting['underline']
                    if formatting['color']:
                        run.font.color.rgb = formatting['color']
                    if formatting['size']:
                        run.font.size = formatting['size']
                    run.font.name = formatting['name']
        else:
            paragraph.add_run(new_text_content)
    
        return 1
    
    def replace_text_in_table_cell(self, cell, old_text, new_text):
        """Обрабатывает все содержимое ячейки таблицы"""
        
        replacements = 0
        # Обрабатываем все параграфы в ячейке
        for paragraph in cell.paragraphs:
            replacements += self.replace_text_in_paragraph(paragraph, old_text, new_text)
        
        # Рекурсивная обработка вложенных таблиц
        for table in cell.tables:
            for row in table.rows:
                for nested_cell in row.cells:
                    replacements += self.replace_text_in_table_cell(nested_cell, old_text, new_text)
        
        return replacements

    def replace_text_in_docx(self, root_folder, old_text, new_text):
        
        total_files = 0
        total_replacements = 0
        
        for foldername, _, filenames in os.walk(root_folder):
            for filename in [f for f in filenames if f.lower().endswith('.docx')]:
                file_path = os.path.join(foldername, filename)
                try:
                    doc = Document(file_path)
                    replacements = 0
                    
                    # Проверка всего документа на наличие текста
                    full_text = '\n'.join(p.text for p in doc.paragraphs)
                    if old_text not in full_text:
                        # Дополнительная проверка таблиц
                        table_text = ''
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    table_text += '\n'.join(p.text for p in cell.paragraphs)
                        if old_text not in table_text:
                            total_files += 1
                            continue
                    
                    # Обработка обычных параграфов
                    for paragraph in doc.paragraphs:
                        replacements += self.replace_text_in_paragraph(paragraph, old_text, new_text)
                    
                    # Улучшенная обработка таблиц
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                replacements += self.replace_text_in_table_cell(cell, old_text, new_text)
                    
                    # Обработка колонтитулов
                    for section in doc.sections:
                        for paragraph in section.header.paragraphs:
                            replacements += self.replace_text_in_paragraph(paragraph, old_text, new_text)
                        for paragraph in section.footer.paragraphs:
                            replacements += self.replace_text_in_paragraph(paragraph, old_text, new_text)
                    
                    if replacements > 0:
                        doc.save(file_path)
                        total_replacements += 1
                    
                    total_files += 1
                    self.status_label.config(text=f"Обработано: {total_files} файлов, замен: {total_replacements}")
                    self.root.update()
                    
                except Exception as e:
                    print(f"Ошибка при обработке {file_path}: {str(e)}")
        
        return total_files, total_replacements


if __name__ == "__main__":
    root = tk.Tk()
    app = DocxReplacerApp(root)
    root.mainloop()